from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
SRC=ROOT / '01-source/original/西安智瞳安宇科技有限公司商业计划书-20260903A.docx'
OUT=ROOT / '08-output/西安智瞳安宇科技有限公司商业计划书-20260903A-Word内部可直接编辑版.docx'
TARGETS={1,2,3,10,11,12,15,16}
NAVY='17365D'; BLUE='2F75B5'; PALE='EAF2F8'; PALE2='F7FAFD'; BORDER='A9C4DC'; TEXT='263645'; MUTED='68798A'; WHITE='FFFFFF'; LIGHT='DCEBF7'
FONT_CN='等线'; FONT_LATIN='Aptos'

from build_bp_docx import replace_text_everywhere, style_tables, style_captions, resize_and_center_figures


def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd');tcPr.append(shd)
    shd.set(qn('w:fill'),fill);shd.set(qn('w:val'),'clear')

def border(cell, **edges):
    tcPr=cell._tc.get_or_add_tcPr(); b=tcPr.first_child_found_in('w:tcBorders')
    if b is None: b=OxmlElement('w:tcBorders');tcPr.append(b)
    for edge, spec in edges.items():
        el=b.find(qn(f'w:{edge}'))
        if el is None: el=OxmlElement(f'w:{edge}');b.append(el)
        color,size,val=spec
        el.set(qn('w:val'),val);el.set(qn('w:sz'),str(size));el.set(qn('w:color'),color);el.set(qn('w:space'),'0')

def set_cell_margins(cell,top=70,start=80,bottom=70,end=80):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar');tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}');tcMar.append(node)
        node.set(qn('w:w'),str(v));node.set(qn('w:type'),'dxa')

def set_run(run,size=9,bold=False,color=TEXT):
    run.font.name=FONT_LATIN;run.font.size=Pt(size);run.font.bold=bold
    rPr=run._r.get_or_add_rPr();rFonts=rPr.rFonts
    if rFonts is None:rFonts=OxmlElement('w:rFonts');rPr.insert(0,rFonts)
    rFonts.set(qn('w:ascii'),FONT_LATIN);rFonts.set(qn('w:hAnsi'),FONT_LATIN);rFonts.set(qn('w:eastAsia'),FONT_CN)
    c=rPr.find(qn('w:color'))
    if c is None:c=OxmlElement('w:color');rPr.append(c)
    c.set(qn('w:val'),color)

def cell_text(cell,text,size=9,bold=False,color=TEXT,align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text='';p=cell.paragraphs[0];p.alignment=align;p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(0);p.paragraph_format.line_spacing=1.0
    lines=str(text).split('\n')
    for i,line in enumerate(lines):
        if i: p.add_run().add_break()
        r=p.add_run(line);set_run(r,size,bold,color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;set_cell_margins(cell)

def tbl_props(table,widths=None):
    table.autofit=False;table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit=False
    if widths:
        for row in table.rows:
            for i,c in enumerate(row.cells):
                c.width=Inches(widths[min(i,len(widths)-1)])
    for row in table.rows:
        trPr=row._tr.get_or_add_trPr();trPr.append(OxmlElement('w:cantSplit'))
        for c in row.cells:
            border(c,top=(BORDER,6,'single'),bottom=(BORDER,6,'single'),left=(BORDER,6,'single'),right=(BORDER,6,'single'))

def exact(row,h):
    row.height=Inches(h);row.height_rule=WD_ROW_HEIGHT_RULE.EXACTLY

def insert_after(doc,p,table):
    p._p.addnext(table._tbl)
    p._p.getparent().remove(p._p)

def add_title_row(table,title,size=12):
    c=table.cell(0,0).merge(table.cell(0,len(table.columns)-1));shade(c,WHITE);cell_text(c,title,size,True,NAVY)
    border(c,bottom=(BLUE,18,'single'),top=(WHITE,0,'nil'),left=(WHITE,0,'nil'),right=(WHITE,0,'nil'))

def make_fig1(doc,p):
    t=doc.add_table(rows=8,cols=3);tbl_props(t,[1.78,1.78,1.78]);add_title_row(t,'三层产品构成同一套企业智能底座',12.5);exact(t.rows[0],.36)
    data=[('组织级','AragonTeam','企业 AI 原生\n人机协同工作站','研发团队与组织管理者'),('通用级','AegisClaw','线上安全通用智能体','跨职能个人与团队'),('行业级','合约智审 LegalLens','大模型智能体驱动的\n智能合同审核系统','法务、采购、合规与项目管理')]
    for ci,(level,name,desc,aud) in enumerate(data):
        for ri,(txt,sz,bold,col) in enumerate([(level,9.3,True,NAVY),(name,10.5,True,NAVY),(desc,8.2,False,TEXT),(aud,7.4,False,MUTED)],start=1):
            c=t.cell(ri,ci);shade(c,WHITE if ri<4 else PALE2);cell_text(c,txt,sz,bold,col)
    for r,h in zip(t.rows[1:5],[.30,.38,.48,.30]):exact(r,h)
    c=t.cell(5,0).merge(t.cell(5,2));shade(c,PALE);cell_text(c,'统一企业智能底座｜四类企业级要件',10.2,True,NAVY);exact(t.rows[5],.34)
    c=t.cell(6,0).merge(t.cell(6,2));shade(c,PALE);cell_text(c,'身份与权限｜智能体是谁、能碰哪些系统与数据    ·    知识与上下文｜企业知识与项目背景可被稳定调用\n流程与状态｜任务有依赖、有状态、可回退可重入    ·    审计与留痕｜谁做、何时做、依据什么，全程可查询',7.3,False,TEXT);exact(t.rows[6],.60)
    c=t.cell(7,0).merge(t.cell(7,2));shade(c,WHITE);cell_text(c,'组织级提供约束与留痕｜通用级提供安全执行力｜行业级提供领域深度',7.6,True,NAVY);exact(t.rows[7],.28)
    insert_after(doc,p,t)

def make_fig2(doc,p):
    pb=doc.add_paragraph(); p._p.addprevious(pb._p); pb.paragraph_format.page_break_before=True; pb.paragraph_format.space_after=Pt(0)
    t=doc.add_table(rows=5,cols=3);tbl_props(t,[1.55,1.55,1.55]);add_title_row(t,'企业级 Agentic AI 市场规模',12.5);exact(t.rows[0],.38)
    labels=[('2025 年','67.6 亿美元',PALE),('五年 CAGR','47%',WHITE),('2030 年（预测）','460.4 亿美元',LIGHT)]
    for i,(h,v,fc) in enumerate(labels):
        c=t.cell(1,i);shade(c,fc);cell_text(c,h,8.5,True,NAVY);c=t.cell(2,i);shade(c,fc);cell_text(c,v,15 if i!=1 else 13,True,BLUE if i==1 else NAVY)
    exact(t.rows[1],.32);exact(t.rows[2],.66)
    c=t.cell(3,0).merge(t.cell(3,2));shade(c,PALE2);cell_text(c,'67.6 → 460.4｜按原 BP 口径，五年复合增长率约 47%',8.5,True,NAVY);exact(t.rows[3],.36)
    c=t.cell(4,0).merge(t.cell(4,2));shade(c,WHITE);cell_text(c,'数据来源：MarketsandMarkets',7.6,False,MUTED);exact(t.rows[4],.28)
    insert_after(doc,p,t)

def make_fig3(doc,p):
    t=doc.add_table(rows=6,cols=4);tbl_props(t,[1.33]*4);add_title_row(t,'企业 AI 能力四级阶梯',12.5);exact(t.rows[0],.38)
    steps=[('L1 对话问答','回答问题','停在个人对话'),('L2 助手辅助','补全与草拟','仍需人工搬运'),('L3 智能体执行','执行多步任务','缺身份 / 权限边界'),('L4 组织智能','组织级协同','进入权限与流程')]
    for i,(h,a,b) in enumerate(steps):
        shade(t.cell(1,i),PALE if i<3 else LIGHT);cell_text(t.cell(1,i),h,9.3,True,NAVY)
        shade(t.cell(2,i),WHITE);cell_text(t.cell(2,i),a,8.2,True,TEXT)
        shade(t.cell(3,i),PALE2);cell_text(t.cell(3,i),b,7.4,False,MUTED)
    for r,h in zip(t.rows[1:4],[.40,.34,.34]):exact(r,h)
    c=t.cell(4,0).merge(t.cell(4,3));shade(c,PALE);cell_text(c,'每上一阶需补齐：身份与权限｜知识与上下文｜流程与状态｜审计与留痕',8.3,True,NAVY);exact(t.rows[4],.44)
    c=t.cell(5,0).merge(t.cell(5,3));shade(c,WHITE);cell_text(c,'前三阶描述“单个智能体能做什么”；L4 描述“组织如何让智能体进入生产流程”',7.6,True,NAVY);exact(t.rows[5],.32)
    insert_after(doc,p,t)

def make_fig10(doc,p):
    t=doc.add_table(rows=6,cols=3);tbl_props(t,[1.78]*3);add_title_row(t,'公司层面的三种交付形态',12.5);exact(t.rows[0],.38)
    cards=[('I 私有化服务器部署','数据、模型、知识与执行\n全部在企业内网闭环','对接现有身份、权限与审计体系','适用：总部级客户'),('II 便携式一体机','本地模型 + AI 算力 + 协作平台\n开箱即用，支持断网运行','面向现场化、临时化部署','适用：项目现场 / 无外网'),('III 私有云服务','私有网络内提供模型、算力\n与协作平台完整能力','兼顾集中管理与私有网络边界','适用：公司与团队办公')]
    for i,(h,b,m,f) in enumerate(cards):
        shade(t.cell(1,i),PALE);cell_text(t.cell(1,i),h,9.1,True,NAVY)
        shade(t.cell(2,i),WHITE);cell_text(t.cell(2,i),b,7.8,True,TEXT)
        shade(t.cell(3,i),PALE2);cell_text(t.cell(3,i),m,7.2,False,MUTED)
        shade(t.cell(4,i),PALE);cell_text(t.cell(4,i),f,7.4,True,NAVY)
    for r,h in zip(t.rows[1:5],[.40,.55,.42,.34]):exact(r,h)
    c=t.cell(5,0).merge(t.cell(5,2));shade(c,WHITE);cell_text(c,'AegisClaw 以线上云端工作空间提供；行业级合约智审采用公有云与私有化双形态',7.4,True,NAVY);exact(t.rows[5],.36)
    insert_after(doc,p,t)

def make_fig11(doc,p):
    t=doc.add_table(rows=6,cols=7);tbl_props(t,[.82,.72,.72,.72,.72,.72,.72]);add_title_row(t,'合约智审系统四层架构',12.5);exact(t.rows[0],.38)
    layers=[('应用层',['智能审查','上下游一致性','资信审查','文稿智审','合同生成','知识库']),('核心服务层',['文档解析','知识抽取','知识图谱','规则引擎','多智能体协同','垂直大模型']),('技术架构层',['前端框架','后端服务','存储与检索','部署与运维']),('基础层',['算力资源','存储资源'])]
    for ri,(lname,items) in enumerate(layers,start=1):
        shade(t.cell(ri,0),PALE);cell_text(t.cell(ri,0),lname,8.6,True,NAVY)
        for ci in range(1,7):
            c=t.cell(ri,ci);shade(c,WHITE if ri%2 else PALE2);cell_text(c,items[ci-1] if ci<=len(items) else '',7.0,ri<3,TEXT)
        exact(t.rows[ri],.46)
    c=t.cell(5,0).merge(t.cell(5,6));shade(c,WHITE);cell_text(c,'基础资源 → 技术架构 → 核心服务 → 应用能力｜四层职责清晰解耦',7.5,True,NAVY);exact(t.rows[5],.32)
    insert_after(doc,p,t)

def make_fig12(doc,p):
    t=doc.add_table(rows=6,cols=5);tbl_props(t,[1.45,.28,1.12,.28,1.45]);add_title_row(t,'收入结构与客户复制路径',12.5);exact(t.rows[0],.38)
    cell_text(t.cell(1,0),'四条收入线',9.0,True,NAVY);cell_text(t.cell(1,2),'标准化实施方法',9.0,True,NAVY);cell_text(t.cell(1,4),'三条复制路径',9.0,True,NAVY)
    for c in [t.cell(1,0),t.cell(1,2),t.cell(1,4)]:shade(c,PALE)
    for ci in [1,3]:shade(t.cell(1,ci),WHITE);cell_text(t.cell(1,ci),'',8)
    exact(t.rows[1],.34)
    left=['私有化部署软件授权','便携式一体机','云服务订阅','行业实施与规则库共建\n及年度运维']; right=['战略渠道复制\n省公司继承已验证口径','行业样板复制\n通信、交通形成量化样板','政企触点\n技术评审与试点立项','']
    center=t.cell(2,2).merge(t.cell(5,2));shade(center,PALE);cell_text(center,'首单共建行业规则库\n与统一审查口径\n\n规则、模板、图谱\n沉淀为领域资产',8.0,True,NAVY)
    for i in range(4):
        r=2+i;shade(t.cell(r,0),WHITE);cell_text(t.cell(r,0),left[i],7.6,True,TEXT);shade(t.cell(r,1),WHITE);cell_text(t.cell(r,1),'→',11,True,BLUE);shade(t.cell(r,3),WHITE);cell_text(t.cell(r,3),'→' if i<3 else '',11,True,BLUE);shade(t.cell(r,4),PALE2 if i<3 else WHITE);cell_text(t.cell(r,4),right[i],7.2,True if i<3 else False,NAVY)
        exact(t.rows[r],.42)
    insert_after(doc,p,t)

def make_fig15(doc,p):
    t=doc.add_table(rows=4,cols=5);tbl_props(t,[1.06]*5);add_title_row(t,'融资资金用途结构（合计 1000 万元）',12.0);exact(t.rows[0],.38)
    seg=[('40%','产品研发与\n技术攻关','400 万元'),('20%','市场与\n渠道建设','200 万元'),('20%','交付与\n实施团队','200 万元'),('12%','一体机产品化\n与备货','120 万元'),('8%','运营储备','80 万元')]
    for i,(pcent,name,amt) in enumerate(seg):
        shade(t.cell(1,i),LIGHT if i==0 else PALE);cell_text(t.cell(1,i),pcent,11,True,NAVY)
        shade(t.cell(2,i),WHITE);cell_text(t.cell(2,i),name,7.5,True,NAVY)
        shade(t.cell(3,i),PALE2);cell_text(t.cell(3,i),amt,8.0,True,BLUE)
    exact(t.rows[1],.42);exact(t.rows[2],.45);exact(t.rows[3],.32)
    insert_after(doc,p,t)

def make_fig16(doc,p):
    t=doc.add_table(rows=4,cols=3);tbl_props(t,[1.78]*3);add_title_row(t,'三年发展里程碑',12.0);exact(t.rows[0],.38)
    headers=['12 个月','24 个月','36 个月'];bullets=[['三层产品商业化版本定版','一体机产品化并完成首批交付','金融行业第二家客户私有化落地'],['规则库覆盖新增两个行业','云订阅形成经常性收入','实施团队支撑并行项目'],['四条收入线全部形成规模','四个行业各建可复用样板','具备向新行业标准化复制能力']]
    for i,h in enumerate(headers):shade(t.cell(1,i),PALE);cell_text(t.cell(1,i),h + ('  →' if i<2 else ''),10.3,True,BLUE)
    exact(t.rows[1],.36)
    for i,lines in enumerate(bullets):shade(t.cell(2,i),WHITE);cell_text(t.cell(2,i),'• ' + '\n• '.join(lines),7.8,False,TEXT,WD_ALIGN_PARAGRAPH.LEFT)
    exact(t.rows[2],.88)
    c=t.cell(3,0).merge(t.cell(3,2));shade(c,PALE2);cell_text(c,'里程碑均按原 BP 内容，不新增未经确认的经营指标',7.2,False,MUTED);exact(t.rows[3],.28)
    insert_after(doc,p,t)

BUILDERS={1:make_fig1,2:make_fig2,3:make_fig3,10:make_fig10,11:make_fig11,12:make_fig12,15:make_fig15,16:make_fig16}

def cleanup_unused_images(docx_path):
    from zipfile import ZipFile, ZIP_DEFLATED
    from lxml import etree
    tmp=docx_path.with_suffix('.clean.tmp.docx')
    RNS='http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    with ZipFile(docx_path,'r') as zin:
        doc_xml=zin.read('word/document.xml'); rel_xml=zin.read('word/_rels/document.xml.rels')
        root=etree.fromstring(doc_xml); used=set(root.xpath('//@r:embed | //@r:link | //@r:id',namespaces={'r':RNS}))
        relroot=etree.fromstring(rel_xml); remove_targets=[]
        for rel in list(relroot):
            rid=rel.get('Id'); target=rel.get('Target',''); typ=rel.get('Type','')
            if '/image' in typ and rid not in used:
                remove_targets.append('word/'+target if not target.startswith('/') else target.lstrip('/')); relroot.remove(rel)
        with ZipFile(tmp,'w',ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in remove_targets: continue
                if item.filename=='word/_rels/document.xml.rels':
                    zout.writestr(item,etree.tostring(relroot,xml_declaration=True,encoding='UTF-8',standalone=True))
                else: zout.writestr(item,zin.read(item.filename))
    tmp.replace(docx_path)

def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    if not SRC.exists(): raise FileNotFoundError(f'Missing source DOCX: {SRC}')
    doc=Document(str(SRC))
    replace_text_everywhere(doc); style_tables(doc); style_captions(doc); resize_and_center_figures(doc)
    counter=0
    for p in list(doc.paragraphs):
        drawings=p._p.xpath('.//*[local-name()="drawing"]')
        if drawings:
            counter += len(drawings)
            if counter in TARGETS: BUILDERS[counter](doc,p)
    if counter!=16: raise RuntimeError(f'Expected 16 images, found {counter}')
    doc.save(str(OUT)); cleanup_unused_images(OUT)
    check=Document(str(OUT))
    if len(check.inline_shapes)!=8: raise RuntimeError(f'Expected 8 remaining product screenshots, found {len(check.inline_shapes)}')
    text='\n'.join(p.text for p in check.paragraphs)
    if 'InkClaw' in text: raise RuntimeError('InkClaw remains')
    print(OUT)

if __name__=='__main__': main()
