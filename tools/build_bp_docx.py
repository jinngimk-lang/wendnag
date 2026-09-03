from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / '01-source/original/西安智瞳安宇科技有限公司商业计划书-20260903A.docx'
OUT = ROOT / '08-output/西安智瞳安宇科技有限公司商业计划书-20260903A-版式优化SVG版.docx'
SVG_DIR = ROOT / '04-assets/figures/svg'
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY='17365D'; BLUE='2F75B5'; PALE='EAF2F8'; PALE2='F7FAFD'; BORDER='C7D7E5'; TEXT='263645'; WHITE='FFFFFF'
FONT_CN='等线'; FONT_LATIN='Aptos'

FIG_SPECS = {
    1:('fig01-three-layer-enterprise-intelligence-base.svg',5.65,3.30),
    2:('fig02-enterprise-agentic-ai-market.svg',5.45,3.45),
    3:('fig03-enterprise-ai-four-level-ladder.svg',5.65,3.20),
    10:('fig10-three-delivery-models.svg',5.65,2.72),
    11:('fig11-legallens-four-layer-architecture.svg',5.65,3.25),
    12:('fig12-revenue-and-customer-replication.svg',5.65,3.15),
    15:('fig15-funding-use-structure.svg',5.65,2.10),
    16:('fig16-three-year-milestones.svg',5.65,2.05),
}

TABLE_WIDTHS = {
    0:[.22,.78], 1:[.39,.18,.17,.26], 2:[.14,.34,.25,.27],
    3:[.14,.13,.19,.20,.20,.14], 4:[.15,.22,.63],
    5:[.15,.17,.18,.20,.30], 6:[.15,.17,.18,.20,.30],
    7:[.18,.22,.14,.25,.21], 8:[.14,.14,.18,.18,.19,.17],
    9:[.15,.12,.22,.17,.17,.17], 10:[.28,.72], 11:[.25,.12,.14,.49]
}

CAPTION_REPL = {
    1:'三层产品与统一企业智能底座',
    2:'企业级 Agentic AI 市场规模',
    3:'企业 AI 能力四级阶梯',
    7:'AegisClaw 多智能体 DAG 协作视图',
    8:'AegisClaw 智能文档协同编辑',
    10:'公司层面的三种交付形态',
}

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None:
        shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill); shd.set(qn('w:val'),'clear')

def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None:
            node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_table_borders(table, color=BORDER, size='6'):
    tblPr=table._tbl.tblPr; borders=tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders=OxmlElement('w:tblBorders'); tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el=borders.find(qn(f'w:{edge}'))
        if el is None:
            el=OxmlElement(f'w:{edge}'); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:space'),'0'); el.set(qn('w:color'),color)

def set_repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def set_cant_split(row):
    trPr=row._tr.get_or_add_trPr(); trPr.append(OxmlElement('w:cantSplit'))

def set_run_font(run, size, bold=False, color=TEXT):
    run.font.name=FONT_LATIN; run.font.size=Pt(size); run.font.bold=bold
    rPr=run._r.get_or_add_rPr(); rFonts=rPr.rFonts
    if rFonts is None:
        rFonts=OxmlElement('w:rFonts'); rPr.insert(0,rFonts)
    rFonts.set(qn('w:ascii'),FONT_LATIN); rFonts.set(qn('w:hAnsi'),FONT_LATIN); rFonts.set(qn('w:eastAsia'),FONT_CN)
    c=rPr.find(qn('w:color'))
    if c is None:
        c=OxmlElement('w:color'); rPr.append(c)
    c.set(qn('w:val'),color)

def replace_text_everywhere(doc):
    replacements=[('AegisClaw（曾用名 InkClaw）','AegisClaw'),('InkClaw','AegisClaw')]
    for p in doc.paragraphs:
        for run in p.runs:
            for a,b in replacements:
                if a in run.text: run.text=run.text.replace(a,b)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for a,b in replacements:
                            if a in run.text: run.text=run.text.replace(a,b)

def style_tables(doc):
    usable = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    for ti,table in enumerate(doc.tables):
        table.autofit=False; table.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_table_borders(table)
        widths=TABLE_WIDTHS.get(ti,[1/len(table.columns)]*len(table.columns))
        body_size=10.0 if len(table.columns)<=2 else (9.0 if len(table.columns)<=4 else 8.0)
        head_size=10.2 if len(table.columns)<=4 else 8.8
        for ri,row in enumerate(table.rows):
            set_cant_split(row)
            if ri==0: set_repeat_header(row)
            for ci,cell in enumerate(row.cells):
                cell.width=int(usable*widths[min(ci,len(widths)-1)])
                cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell,top=65,bottom=65,start=70,end=70)
                if ri==0: set_cell_shading(cell,BLUE)
                elif ci==0: set_cell_shading(cell,PALE)
                elif ri%2==0: set_cell_shading(cell,PALE2)
                else: set_cell_shading(cell,WHITE)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER if ri==0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        set_run_font(run,head_size if ri==0 else body_size,bold=(ri==0 or ci==0),color=WHITE if ri==0 else (NAVY if ci==0 else TEXT))

def restyle_caption(p, kind, num, title_text):
    p.clear(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.0
    if kind=='表': p.paragraph_format.keep_with_next=True
    r=p.add_run(f'{kind} {num}  '); set_run_font(r,9.5,True,NAVY)
    r=p.add_run(title_text); set_run_font(r,9.5,False,TEXT)

def style_captions(doc):
    fig_pat=re.compile(r'^图\s*(\d+)\s*[　\s]*(.*)$'); tbl_pat=re.compile(r'^表\s*(\d+)\s*[　\s]*(.*)$')
    for p in doc.paragraphs:
        t=p.text.strip(); m=fig_pat.match(t)
        if m:
            n=int(m.group(1)); title=CAPTION_REPL.get(n,m.group(2).strip()).replace('InkClaw','AegisClaw')
            restyle_caption(p,'图',n,title); continue
        m=tbl_pat.match(t)
        if m: restyle_caption(p,'表',int(m.group(1)),m.group(2).strip())

def resize_and_center_figures(doc):
    for idx,shape in enumerate(doc.inline_shapes, start=1):
        p=shape._inline.getparent().getparent().getparent(); pPr=p.get_or_add_pPr(); jc=pPr.find(qn('w:jc'))
        if jc is None: jc=OxmlElement('w:jc'); pPr.append(jc)
        jc.set(qn('w:val'),'center')
        if idx in FIG_SPECS:
            _,w,h=FIG_SPECS[idx]; shape.width=Inches(w); shape.height=Inches(h)
        else:
            oldw,oldh=shape.width,shape.height; neww=Inches(5.25); shape.width=neww; shape.height=int(oldh*(neww/oldw))

def patch_svg_images(docx_path):
    mapping={'image1.png':FIG_SPECS[1][0],'image2.png':FIG_SPECS[2][0],'image3.png':FIG_SPECS[3][0],'image10.png':FIG_SPECS[10][0],'image11.png':FIG_SPECS[11][0],'image12.png':FIG_SPECS[12][0],'image15.png':FIG_SPECS[15][0],'image16.png':FIG_SPECS[16][0]}
    tmp=docx_path.with_suffix('.tmp.docx')
    with ZipFile(docx_path,'r') as zin, ZipFile(tmp,'w',ZIP_DEFLATED) as zout:
        rels=zin.read('word/_rels/document.xml.rels').decode('utf-8'); cts=zin.read('[Content_Types].xml').decode('utf-8')
        if 'Extension="svg"' not in cts: cts=cts.replace('</Types>','<Default Extension="svg" ContentType="image/svg+xml"/></Types>')
        for old,new in mapping.items(): rels=rels.replace(f'Target="media/{old}"',f'Target="media/{new}"')
        for item in zin.infolist():
            if item.filename=='word/_rels/document.xml.rels': zout.writestr(item,rels.encode('utf-8'))
            elif item.filename=='[Content_Types].xml': zout.writestr(item,cts.encode('utf-8'))
            elif item.filename.startswith('word/media/') and Path(item.filename).name in mapping: continue
            else: zout.writestr(item,zin.read(item.filename))
        for new in mapping.values(): zout.writestr(f'word/media/{new}',(SVG_DIR/new).read_bytes())
    tmp.replace(docx_path)

def main():
    if not SRC.exists(): raise FileNotFoundError(f'Missing source DOCX: {SRC}')
    doc=Document(str(SRC)); replace_text_everywhere(doc); style_tables(doc); style_captions(doc); resize_and_center_figures(doc); doc.save(str(OUT)); patch_svg_images(OUT)
    with ZipFile(OUT,'r') as z:
        xml=z.read('word/document.xml').decode('utf-8'); assert 'InkClaw' not in xml
        names=set(z.namelist())
        for svg,_,_ in FIG_SPECS.values(): assert f'word/media/{svg}' in names
    print(OUT)

if __name__=='__main__': main()
